import os
import tempfile
from pathlib import Path
from typing import Optional, Tuple
import markdown
import base64


class ResumeExporter:
    """Handle resume export in different formats"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
    
    def export_markdown(self, resume_content: str, filename: str = "resume.md") -> Tuple[str, str]:
        """Export resume as Markdown file"""
        file_path = os.path.join(self.temp_dir, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(resume_content)
        return file_path, resume_content
    
    def export_html(self, resume_content: str, filename: str = "resume.html") -> Tuple[str, str]:
        """Export resume as HTML file"""
        # Convert markdown to HTML
        html_content = markdown.markdown(resume_content, extensions=['markdown.extensions.tables'])
        
        # Add basic styling
        styled_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Resume</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        h3 {{
            color: #7f8c8d;
        }}
        ul {{
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 5px;
        }}
        strong {{
            color: #2c3e50;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        
        file_path = os.path.join(self.temp_dir, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(styled_html)
        return file_path, styled_html
    
    def export_pdf(self, resume_content: str, filename: str = "resume.pdf") -> Optional[Tuple[str, bytes]]:
        """Export resume as PDF file (requires additional dependencies)"""
        try:
            # First convert to HTML
            html_path, html_content = self.export_html(resume_content, "temp_resume.html")
            
            # Try to use weasyprint for PDF conversion
            try:
                import weasyprint
                pdf_path = os.path.join(self.temp_dir, filename)
                weasyprint.HTML(html_path).write_pdf(pdf_path)
                
                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_path, pdf_content
            except ImportError:
                # Fallback: try using pdfkit (requires wkhtmltopdf)
                try:
                    import pdfkit
                    pdf_path = os.path.join(self.temp_dir, filename)
                    pdfkit.from_file(html_path, pdf_path)
                    
                    with open(pdf_path, 'rb') as f:
                        pdf_content = f.read()
                    return pdf_path, pdf_content
                except (ImportError, OSError):
                    return None
        except Exception as e:
            print(f"PDF export failed: {e}")
            return None
    
    def export_docx(self, resume_content: str, filename: str = "resume.docx") -> Optional[Tuple[str, bytes]]:
        """Export resume as DOCX file"""
        try:
            from docx import Document
            from docx.shared import Inches
            import re
            
            doc = Document()
            
            # Parse markdown content and convert to DOCX
            lines = resume_content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Main heading (# Title)
                if line.startswith('# '):
                    title = line[2:].strip()
                    heading = doc.add_heading(title, level=1)
                    heading.bold = True
                
                # Section headings (## Section)
                elif line.startswith('## '):
                    section = line[3:].strip()
                    doc.add_heading(section, level=2)
                
                # Subsection headings (### Subsection)
                elif line.startswith('### '):
                    subsection = line[4:].strip()
                    doc.add_heading(subsection, level=3)
                
                # Bullet points
                elif line.startswith('- '):
                    bullet_text = line[2:].strip()
                    # Remove markdown formatting
                    bullet_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)  # Bold
                    bullet_text = re.sub(r'\*(.*?)\*', r'\1', bullet_text)      # Italic
                    doc.add_paragraph(bullet_text, style='List Bullet')
                
                # Regular paragraphs
                else:
                    # Remove markdown formatting
                    clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
                    clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)  # Italic
                    doc.add_paragraph(clean_line)
            
            # Save the document
            docx_path = os.path.join(self.temp_dir, filename)
            doc.save(docx_path)
            
            with open(docx_path, 'rb') as f:
                docx_content = f.read()
            return docx_path, docx_content
            
        except ImportError:
            print("python-docx not installed. Cannot export to DOCX format.")
            return None
        except Exception as e:
            print(f"DOCX export failed: {e}")
            return None
    
    def get_download_link(self, file_path: str, filename: str) -> str:
        """Generate a download link for Streamlit"""
        try:
            with open(file_path, 'rb') as f:
                data = f.read()
            
            # Encode file for download
            b64 = base64.b64encode(data).decode()
            
            # Determine MIME type
            if filename.endswith('.pdf'):
                mime_type = 'application/pdf'
            elif filename.endswith('.docx'):
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif filename.endswith('.html'):
                mime_type = 'text/html'
            else:
                mime_type = 'text/plain'
            
            return f'<a href="data:{mime_type};base64,{b64}" download="{filename}">Download {filename}</a>'
        except Exception as e:
            print(f"Failed to create download link: {e}")
            return ""
    
    def cleanup(self):
        """Clean up temporary files"""
        import shutil
        try:
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Failed to cleanup temp directory: {e}")
    
    def __del__(self):
        """Cleanup on object destruction"""
        self.cleanup()


# Global exporter instance
_exporter = None

def get_exporter() -> ResumeExporter:
    """Get global exporter instance"""
    global _exporter
    if _exporter is None:
        _exporter = ResumeExporter()
    return _exporter